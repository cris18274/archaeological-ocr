"""
ArcheoOCR Backend - IA Ensemble v3
"""
import os
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
os.environ["PYTHONUTF8"] = "1"
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import (JSONResponse, FileResponse,
                                RedirectResponse, StreamingResponse)
from fastapi.staticfiles import StaticFiles
import os, re, io, uuid, shutil, hashlib, json, asyncio, time, difflib
from typing import Dict, Any, List
import easyocr
import cv2
import numpy as np
import pandas as pd
from paddleocr import PaddleOCR, PPStructure
from pdf2image import convert_from_path

# --- CONFIGURACIÓN ESTRATEGIA ARQUEOLÓGICA (8 REGLAS) ---
CONF_THRESHOLD  = 0.45    # Regla 2.2: Descartar si conf < 0.45 (legacy)
CONF_THRESHOLD_V2 = 0.85  # Mejora 1 (ocr2.py v2): Umbral elevado para recognize_cells_batch
INNER_PADDING   = 3       # Regla 1.1: Recorte interno para evitar bordes
MIN_CELL_HEIGHT = 60      # Regla 1.3: Upscaling a min 60px
DATA_ALLOWLIST  = "0123456789 ABCDEFGHIJK" # Regla 2.1

# Vocabulario arqueológico para fuzzy matching (Regla 3.3)
ARCHAEO_VOCAB = [
    "CATALOGO", "VARIEDAD", "TIESTO", "CUERPO", "ALCANCE", "PUNTA", "TOTAL", 
    "CERAMICA", "MAIZ", "BOTELLA", "OLLA", "CUENCO", "PLATO", "JARRA", "COMPOTERA",
    "DECORADO", "INCISO", "PINTADO", "ROJO", "NEGRO", "BLANCO", "CAFÉ", "GRIS",
    "LIMPIEZA", "RECOLECCION", "EXCAVACION", "NIVEL", "UNIDAD", "CUADRO", "LOTE",
    "A Panzaleo", "Adorno", "Adornos tallados", "VASO", "TRIPODE", "ASA", "REBORDE"
]

# Vocabulario extendido Página 3 (Mejora 4 — ocr2.py v2)
VOCAB_ARQUEOLOGICO_V2 = {
    "Lamina":                  "Lámina",
    "Nodulo":                  "Nódulo",
    "Nucleo":                  "Núcleo",
    "Apendice":                "Apéndice",
    "Ceramica":                "CERAMICA",
    "Ceramica total":          "CERAMICA TOTAL",
    "Litica tallada":          "LITICA TALLADA",
    "Percutor Yunque":         "Percutor / Yunque",
    "Reutilizador Torrero":    "Reutilizador / Torrero",
    "Reutilizado":             "Reutilizador / Torrero",
    "Pedestal fragmento":      "Pedestal / Fragmento",
    "Base fragmento":          "Base / Fragmento",
    "Cuerpo deco sin pc":      "Cuerpo deco sin PC",
    "Cuerpo deco con pc":      "Cuerpo deco con PC",
    "Cuerpo no deco sin pc":   "Cuerpo no deco sin PC",
    "Cuerpo no deco con pc":   "Cuerpo no deco con PC",
    "Vasija reconstruida":     "Vasija reconstruida",
    "Vasija completa":         "Vasija completa",
    "Indefinido":              "Indefinido",
    "Indefinidos":             "Indefinidos",
    "Figurilla fragmentada":   "Figurilla fragmentada",
    "Instrumento musical":     "Instrumento musical",
    "Borde decorado -10":      "Borde decorado (-10%)",
    "Borde decorado 10":       "Borde decorado (+10%)",
}

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[WARN] python-docx no instalado — export Word no disponible")

try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False
    print("[WARN] fpdf2 no instalado — export PDF no disponible")


# --------------------------------------------------------------------------
# App y CORS
# --------------------------------------------------------------------------
app = FastAPI(title="ArcheoOCR Advanced")

@app.get("/")
async def root():
    return RedirectResponse(url="/static/index.html")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# --------------------------------------------------------------------------
# Directorios
# --------------------------------------------------------------------------
BASE_DIR     = os.path.dirname(os.path.abspath(__file__))
UPLOADS_DIR  = os.path.normpath(os.path.join(BASE_DIR, "..", "uploads"))
OUTPUT_DIR   = os.path.normpath(os.path.join(BASE_DIR, "..", "output"))
POPPLER_DIR  = os.path.normpath(os.path.join(BASE_DIR, "..", "..", "poppler",
               "poppler-24.08.0", "Library", "bin"))

os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR,  exist_ok=True)

# DLL paths (Windows CUDA/MKL/cuDNN/cuBLAS)
# IMPORTANTE: NVIDIA libs van al FINAL del PATH para no sobrescribir los
# MKL (Math Kernel Library) de PaddlePaddle que estan en el PATH primero.
_NVIDIA_PKGS = r"C:\Users\Estudiantes\AppData\Local\Programs\Python\Python312\Lib\site-packages\nvidia"
for dll_path in [
    r"C:\Program Files\SPGlobal\KingdomSuite\TKS 2023",
    os.path.normpath(os.path.join(BASE_DIR, "..", "..", "venv", "Library", "bin")),
    os.path.join(_NVIDIA_PKGS, "cudnn",  "bin"),   # cudnn64_8/9.dll
    os.path.join(_NVIDIA_PKGS, "cublas", "bin"),   # cublasLt64_12.dll
    os.path.join(_NVIDIA_PKGS, "cuda_runtime", "bin"),  # cudart64.dll (si existe)
]:
    if os.path.exists(dll_path):
        # Agregar al FINAL del PATH (no al inicio) para no sobrescribir MKL de Paddle
        os.environ["PATH"] = os.environ.get("PATH", "") + os.pathsep + dll_path
        if hasattr(os, "add_dll_directory"):
            try:
                os.add_dll_directory(dll_path)
                print(f"DLL: {dll_path}")
            except Exception:
                pass

# Servir archivos estáticos
FRONTEND_DIR = os.path.normpath(os.path.join(BASE_DIR, "..", "frontend"))
app.mount("/static",  StaticFiles(directory=FRONTEND_DIR, html=True), name="static")
app.mount("/uploads", StaticFiles(directory=UPLOADS_DIR),              name="uploads")

print(f"[CONFIG] UPLOADS_DIR  = {UPLOADS_DIR}")
print(f"[CONFIG] FRONTEND_DIR = {FRONTEND_DIR}")

# --------------------------------------------------------------------------
# CONFIGURACIÓN DE PADDLE (GPU ACELERADO)
# --------------------------------------------------------------------------
USE_GPU = True
PADDLE_USE_GPU = True 
print(f"[*] Modo de ejecución: CUDA/GPU (Acelerado)")


# --------------------------------------------------------------------------
# Motores OCR (LAZY LOADED)
# --------------------------------------------------------------------------
ocr_engine      = None
easyocr_engine  = None
table_engine    = None

# Limpiar caché interno de PaddleOCR al arrancar
_paddle_cache = os.path.join(UPLOADS_DIR, "_cache")
if os.path.exists(_paddle_cache):
    try: shutil.rmtree(_paddle_cache); print(f"  [cache] Limpiado: {_paddle_cache}")
    except: pass

def get_easyocr():
    global easyocr_engine
    if easyocr_engine is None:
        print("  [Init] Inicializando EasyOCR (IA Engine 2) en GPU...")
        try:
            import easyocr
            easyocr_engine = easyocr.Reader(['es'], gpu=True)
            print("  [Init] EasyOCR IA Engine inicializado en GPU con éxito.")
        except Exception as e:
            print(f"  [Init] Error GPU: {e}. Reintentando en CPU...")
            import easyocr
            easyocr_engine = easyocr.Reader(['es'], gpu=False)
            print("  [Init] EasyOCR IA Engine inicializado en CPU.")
    return easyocr_engine

def get_table_engine():
    global table_engine
    if table_engine is None:
        print("  [Init] Inicializando PPStructure (Layout Analysis)...")
        from paddleocr import PPStructure
        table_engine = PPStructure(show_log=False, image_orientation=True, use_gpu=True)
        print("  [Init] PPStructure V1 inicializado correctamente.")
    return table_engine

def get_paddle_ocr():
    global ocr_engine
    if ocr_engine is None:
        print("  [Init] Inicializando PaddleOCR (Text Recognition)...")
        from paddleocr import PaddleOCR
        ocr_engine = PaddleOCR(use_angle_cls=True, lang='es', show_log=False, use_gpu=True)
        print("  [Init] PaddleOCR (ES) inicializado correctamente.")
    return ocr_engine

# Engines are now lazy-loaded via get_easyocr(), get_table_engine(), get_paddle_ocr()


# -----------------------------------------------------------------
def imread_unicode(path):
    """Leer archivos con caracteres especiales en Windows (ocr2.py style)"""
    try:
        return cv2.imdecode(np.fromfile(path, dtype=np.uint8), cv2.IMREAD_COLOR)
    except Exception as e:
        print(f"  [imread] Error: {e}")
        return None

def imwrite_unicode(path, img):
    """Guardar archivos con caracteres especiales (ocr2.py style)"""
    try:
        ext = os.path.splitext(path)[1]
        if not ext: ext = ".jpg"
        is_success, buffer = cv2.imencode(ext, img)
        if is_success:
            buffer.tofile(path)
            return True
        return False
    except Exception as e:
        print(f"  [imwrite] Error: {e}")
        return False

def save_step(img_arr, page_id: str, step_num: int) -> str:
    """Guarda imagen de proceso y retorna URL relativa /uploads/..."""
    filename = f"step_{page_id}_{step_num}.jpg"
    path     = os.path.join(UPLOADS_DIR, filename)
    imwrite_unicode(path, img_arr)
    return f"/uploads/{filename}"

# --------------------------------------------------------------------------
# CACHE DE RESULTADOS - Evita reprocesar el mismo archivo
# --------------------------------------------------------------------------
CACHE_DIR = os.path.join(UPLOADS_DIR, "_cache")
os.makedirs(CACHE_DIR, exist_ok=True)

def file_hash(path: str) -> str:
    """Calcula hash MD5 del archivo para usar como clave de caché."""
    h = hashlib.md5()
    h.update(b"v2_batch_gpu") # Salt para invalidar cache vieja y forzar el nuevo motor
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()

def get_cached(file_md5: str) -> dict | None:
    """Retorna resultado cacheado si existe, None si no."""
    cache_path = os.path.join(CACHE_DIR, f"{file_md5}.json")
    if os.path.exists(cache_path):
        try:
            with open(cache_path, "r", encoding="utf-8") as f:
                print(f"  [cache] HIT {file_md5[:8]}...")
                return json.load(f)
        except Exception:
            pass
    return None

def set_cached(file_md5: str, data: dict) -> None:
    """Guarda resultado en caché."""
    cache_path = os.path.join(CACHE_DIR, f"{file_md5}.json")
    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
        print(f"  [cache] SAVE {file_md5[:8]}...")
    except Exception as e:
        print(f"  [cache] Error guardando: {e}")


# --------------------------------------------------------------------------
# SSE PROGRESS - Estado de jobs para progreso en tiempo real
# --------------------------------------------------------------------------
JOB_PROGRESS: Dict[str, Dict[str, Any]] = {}

def update_progress(job_id: str, pct: int, msg: str, done: bool = False) -> None:
    """Actualiza el estado de progreso de un job."""
    JOB_PROGRESS[job_id] = {
        "pct": pct,
        "msg": msg,
        "done": done,
        "ts":   time.time()
    }


# --------------------------------------------------------------------------
# REFUERZO DE BORDES - Fortalece líneas de tabla tenues
# --------------------------------------------------------------------------

def enhance_table_borders(gray_img):
    """Refuerza bordes de tabla débiles usando morfología.
    Aplica dilatación controlada en dirección H y V para conectar líneas rotas."""
    # Kernels para reforzar líneas horizontales y verticales por separado
    h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
    v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 25))

    # Extraer líneas H y V con morfología abrir
    h_lines = cv2.morphologyEx(gray_img, cv2.MORPH_OPEN, h_kernel)
    v_lines = cv2.morphologyEx(gray_img, cv2.MORPH_OPEN, v_kernel)

    # Combinar en una imagen de bordes reforzados
    combined = cv2.add(h_lines, v_lines)

    # Dilatar ligeramente para conectar segmentos rotos
    dilate_k = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
    enhanced  = cv2.dilate(combined, dilate_k, iterations=2)
    return enhanced


# --------------------------------------------------------------------------
# DETECCIÓN DE TABLAS SIN BORDES - Proyección de histogramas
# --------------------------------------------------------------------------

def detect_borderless_table(img):
    """Detecta tablas sin líneas visibles usando proyección de histogramas de
    densidad de texto. Retorna lista de filas de celdas aproximadas (BBOXes)."""
    try:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY) if len(img.shape) == 3 else img
        # Binarizar: texto = blanco
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)

        # Proyección horizontal: densidad de píxeles por fila
        h_proj = np.sum(binary, axis=1).astype(np.float32)
        h, w   = binary.shape

        # Suavizar para encontrar bandas de texto (ventana deslizante)
        smooth_h = np.convolve(h_proj, np.ones(5) / 5, mode='same')

        # Umbral = 5% del máximo
        thresh  = float(np.max(smooth_h)) * 0.05
        in_row  = False
        row_start = 0
        rows    = []

        for y, val in enumerate(smooth_h):
            if val > thresh and not in_row:
                in_row    = True
                row_start = y
            elif val <= thresh and in_row:
                in_row = False
                h_span = y - row_start
                if h_span > 8:
                    rows.append((row_start, y))

        # Proyección vertical para encontrar columnas
        v_proj   = np.sum(binary, axis=0).astype(np.float32)
        smooth_v = np.convolve(v_proj, np.ones(5) / 5, mode='same')
        thresh_v = float(np.max(smooth_v)) * 0.02
        in_col   = False
        col_start = 0
        cols     = []

        for x, val in enumerate(smooth_v):
            if val > thresh_v and not in_col:
                in_col    = True
                col_start = x
            elif val <= thresh_v and in_col:
                in_col = False
                w_span = x - col_start
                if w_span > 15:
                    cols.append((col_start, x))

        if len(rows) < 2 or len(cols) < 2:
            return []   # No hay estructura tabular clara

        # Construir celdas como intersecciones fila × columna
        cells = []
        for r_start, r_end in rows:
            row_cells = []
            for c_start, c_end in cols:
                row_cells.append((c_start, r_start, c_end, r_end))
            cells.append(row_cells)

        print(f"  [borderless] {len(rows)} filas x {len(cols)} cols detectadas por proyeccion")
        return cells

    except Exception as e:
        print(f"  [borderless] Error: {e}")
        return []


# --------------------------------------------------------------------------
# OCR CON FILTRADO POR CONFIANZA
# --------------------------------------------------------------------------
CONF_THRESHOLD = 0.40   # Umbral de confianza reducido: espanol con acentos da <60%

def run_ocr_with_confidence(cell_img, threshold: float = CONF_THRESHOLD):
    """Corre OCR y retorna solo líneas que superan el umbral de confianza.
    Retorna (texto_completo: str, confianza_prom: float, lineas: list)"""
    if get_paddle_ocr() is None:
        return "", 0.0, []
    try:
        result = get_paddle_ocr().ocr(cell_img, cls=True)
        if not result or not result[0]:
            return "", 0.0, []

        lines     = result[0]
        accepted  = []
        skipped   = 0

        for ln in lines:
            text = ln[1][0]
            conf = float(ln[1][1]) if len(ln[1]) > 1 else 1.0
            if conf >= threshold:
                accepted.append((text, conf))
            else:
                skipped += 1

        if skipped:
            print(f"    [conf] {skipped} lineas descartadas (conf<{threshold:.0%})")

        if not accepted:
            return "", 0.0, []

        texts  = [t for t, _ in accepted]
        confs  = [c for _, c in accepted]
        return " ".join(texts), float(np.mean(confs)), accepted

    except Exception as e:
        print(f"  [OCR confidence] Error: {e}")
        return "", 0.0, []

# --------------------------------------------------------------------------
# MEJORAS DE PRECISIÓN - Auto Upscaling, Deskewing y Post-procesamiento
# --------------------------------------------------------------------------

def upscale_if_needed(img, min_side: int = 1500):
    """Escala la imagen al doble si algún lado es menor a min_side píxeles.
    Imágenes pequeñas/de baja resolución producen texto borroso que falla en OCR.
    Solo escala hacia arriba, nunca reduce."""
    if img is None or img.size == 0:
        return img
    h, w = img.shape[:2]
    min_dim = min(h, w)
    if min_dim < min_side:
        scale = min_side / min_dim
        new_w = int(w * scale)
        new_h = int(h * scale)
        img = cv2.resize(img, (new_w, new_h), interpolation=cv2.INTER_CUBIC)
        print(f"  [upscale] {w}x{h} -> {new_w}x{new_h} (x{scale:.2f})")
    return img


def deskew_image(img):
    """Detecta y corrige inclinación de hasta ±30° en la imagen completa.
    Usa Canny + HoughLines para encontrar ángulo dominante de las líneas."""
    try:
        gray    = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY) if len(img.shape) == 3 else img.copy()
        edges   = cv2.Canny(gray, 50, 150, apertureSize=3)
        lines   = cv2.HoughLines(edges, 1, np.pi / 180, threshold=min(gray.shape[:2]) // 4)

        if lines is None or len(lines) < 3:
            return img  # No hay suficientes líneas para estimar ángulo

        # Calcular ángulos (ignorar líneas casi verticales que distorsionan)
        angles = []
        for ln in lines:
            theta = float(ln[0][1])
            deg   = np.degrees(theta) - 90
            if abs(deg) < 30:    # Solo corregir inclinaciones pequeñas-medianas
                angles.append(deg)

        if not angles:
            return img

        angle = float(np.median(angles))
        if abs(angle) < 0.5:     # Menos de 0.5° - no vale la pena rotar
            return img

        h, w = img.shape[:2]
        M    = cv2.getRotationMatrix2D((w / 2, h / 2), angle, 1.0)
        deskewed = cv2.warpAffine(img, M, (w, h),
                                  flags=cv2.INTER_CUBIC,
                                  borderMode=cv2.BORDER_REPLICATE)
        print(f"  [deskew] Corregidos {angle:.2f} grados")
        return deskewed
    except Exception as e:
        print(f"  [deskew] Error: {e}")
        return img



def postprocess_cell_text(text: str) -> str:
    """
    Limpieza y normalizacion de celdas.
    Prioriza vocabulario arqueologico para evitar confusiones de digitos en etiquetas.
    """
    if not text: return ""
    text = text.strip().strip("|-_./")
    
    # 1. Fuzzy match contra vocabulario Arqueológico (Critico para cabeceras)
    if len(text) > 3:
        matches = difflib.get_close_matches(text.upper(), ARCHAEO_VOCAB, n=1, cutoff=0.7)
        if matches: return matches[0]
            
    return text

def clean_ocr_text(text: str) -> str:
    """
    Sincronizado con ocr.py: Filtra ruido si la densidad de caracteres validos es baja (<10%).
    """
    if not text: return ""
    text = re.sub(r'\s+', ' ', text.strip())
    # Regla ocr.py: Si el texto tiene menos de 10% de caracteres alfanuméricos/ES, se considera ruido.
    valid_chars = re.sub(r'[^a-zA-Z0-9áéíóúÁÉÍÓÚüÜñÑ.,-]', '', text)
    if len(valid_chars) < len(text) * 0.1:
        return ""
    return text

# --------------------------------------------------------------------------
# PREPROCESADO DE CELDA  (de ocr.py)
# --------------------------------------------------------------------------

# --- CONFIGURACIÓN ESTRATEGIA ULTRA-MICRO (Baja Resolución) ---
CONF_THRESHOLD = 0.1  
OUTER_MARGIN = 2      
BIN_ADAPTIVE_C = 7    
DATA_ALLOWLIST = "0123456789 -_./ABCDEFGHIJKabcdefghijk"

def _run_ocr_on_image_with_conf(img_arr, cell_id_suffix: str, det=False):
    """Fallback individual para celdas cuando batch falla."""
    if img_arr is None or img_arr.size == 0 or easyocr_engine is None:
        return "", 0.0
    try:
        res = get_easyocr().readtext(img_arr, detail=1) # Usar detail=1 para conf
        if res:
            best = max(res, key=lambda x: x[2])
            return best[1], float(best[2])
        return "", 0.0
    except:
        return "", 0.0

def preprocess_page_forensic(img):
    """Mejora visual de la página para facilitar la segmentación de tablas."""
    if img is None or img.size == 0: return img
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    # CLAHE para contraste
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    enhanced = clahe.apply(gray)
    return cv2.cvtColor(enhanced, cv2.COLOR_GRAY2BGR)

def preprocess_cell_clahe(img):
    if img is None or img.size == 0: return img
    if len(img.shape) == 3:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    else: gray = img
    # Upscaling (simulando x2 de ocr.py)
    h, w = gray.shape[:2]
    gray = cv2.resize(gray, (w*2, h*2), interpolation=cv2.INTER_CUBIC)
    # CLAHE para contraste local (clave en ocr.py)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    enhanced = clahe.apply(gray)
    return cv2.cvtColor(enhanced, cv2.COLOR_GRAY2BGR)

def detect_table_morphological(img):
    """
    Estrategia de segmentación por morfología (líneas), útil para Page 3.
    """
    if img is None or img.size == 0: return []
    
    # 1. Red channel for best grid contrast (Cyan/Blue grid lines)
    if len(img.shape) == 3:
        img_gray = img[:,:,2] # Red channel
    else: img_gray = img
    
    # 2. Adaptive Thresholding for grid detection
    img_bin = cv2.adaptiveThreshold(img_gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                   cv2.THRESH_BINARY_INV, 15, 10)

    # REGLA ocr.py: Usa el ancho (shape[1]) para ambos kernels
    W_ref = img_gray.shape[1]
    
    # Líneas Verticales
    kl_v = max(W_ref // 120, 3)
    v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, kl_v))
    v_lines = cv2.erode(img_bin, v_kernel, iterations=3)
    v_lines = cv2.dilate(v_lines, v_kernel, iterations=3)

    # Líneas Horizontales
    kl_h = max(W_ref // 40, 3)
    h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (kl_h, 1))
    h_lines = cv2.erode(img_bin, h_kernel, iterations=3)
    h_lines = cv2.dilate(h_lines, h_kernel, iterations=3)

    # Grid
    grid = cv2.addWeighted(v_lines, 0.5, h_lines, 0.5, 0.0)
    grid = cv2.erode(cv2.bitwise_not(grid), cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3)), iterations=2)
    _, grid = cv2.threshold(grid, 0, 255, cv2.THRESH_OTSU)

    contours, _ = cv2.findContours(grid, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
    boxes = []
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        # REGLA ocr.py: Filtro de area mas restringido (>400) para evitar ruido/letras
        if cv2.contourArea(cnt) > 400 and 10 < h < 250:
            boxes.append((x, y, w, h))
    return boxes

def get_skew_angle(img_path):
    """Detecta el ángulo de inclinación del texto (vía ocr.py)"""
    try:
        raw = get_paddle_ocr().ocr(img_path, cls=True)
        if raw and raw[0]:
            angles = [line[1][2] for line in raw[0] if len(line[1]) > 2]
            if angles: return np.mean(angles)
    except: pass
    return 0

def preprocess_cell(cell_img):
    """
    Preprocesado MEJORADO v2 (ocr2.py mejoras 2+3):
    - Mejora 3: Pre-rotación si la celda es más alta que ancha (texto vertical rotado 90°)
    - Mejora 2: Zoom adaptativo 2x-4x según alto de celda
    - CLAHE para contraste local
    """
    if cell_img is None or cell_img.size == 0:
        return cell_img

    if len(cell_img.shape) == 3:
        gray = cv2.cvtColor(cell_img, cv2.COLOR_BGR2GRAY)
    else:
        gray = cell_img.copy()

    raw_h, raw_w = gray.shape[:2]

    # ── Mejora 3: Pre-rotación para celdas con etiquetas verticales ──────────
    # En Page 3, las etiquetas de filas están rotadas 90° (h >> w)
    if raw_h > raw_w * 1.5:
        gray   = cv2.rotate(gray, cv2.ROTATE_90_CLOCKWISE)
        cell_h = gray.shape[0]
    else:
        cell_h = raw_h

    # ── Mejora 2: Zoom adaptativo según tamaño ───────────────────────────────
    if cell_h < 20:
        zoom = 4
    elif cell_h < 30:
        zoom = 3
    elif cell_h < 45:
        zoom = 3
    else:
        zoom = 2

    resized  = cv2.resize(gray, None, fx=zoom, fy=zoom, interpolation=cv2.INTER_CUBIC)
    blurred  = cv2.GaussianBlur(resized, (3, 3), 0)
    clahe    = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(blurred)
    return enhanced


def _post_process_v2(text: str) -> str:
    """Mejora 4: Normaliza terminología usand vocabulario arqueológico extendido."""
    t = text.strip()
    if not t or len(t) < 2:
        return t
    best_score = 0.0
    best_match = t
    for wrong, correct in VOCAB_ARQUEOLOGICO_V2.items():
        score = difflib.SequenceMatcher(None, t.lower(), wrong.lower()).ratio()
        if score > best_score and score >= 0.88:
            best_score = score
            best_match = correct
    return best_match


def recognize_cells_batch(full_img, boxes, rotate_deg=0, allowlist=None, margin=5):
    """
    Pipeline OCR mejorado v2 — sincronizado con ocr2.py (5 mejoras):
      Mejora 1: Filtro de confianza 0.85
      Mejora 2: Zoom adaptativo (2x-4x) según alto de celda
      Mejora 3: Pre-rotación de celdas con texto vertical
      Mejora 4: Post-proceso con vocabulario arqueológico
      Mejora 5: Margen proporcional al alto de celda
    """
    if not boxes or full_img is None:
        return []

    paddle_engine = get_paddle_ocr()
    clahe         = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    final_cells   = []

    for i, box in enumerate(boxes):
        if len(box) != 4:
            final_cells.append("")
            continue

        x, y, w, h = box

        # ── Mejora 5: Margen proporcional al alto de celda ───────────────────
        margin_px = max(3, min(8, h // 5))
        iy  = max(y - margin_px, 0)
        ix  = max(x - margin_px, 0)
        roi = full_img[iy : y + h + margin_px, ix : x + w + margin_px]

        if roi.size == 0:
            final_cells.append("")
            continue

        gray_roi = cv2.cvtColor(roi, cv2.COLOR_BGR2GRAY) if len(roi.shape) == 3 else roi.copy()
        roi_h, roi_w = gray_roi.shape[:2]

        # ── Mejora 3: Pre-rotación para etiquetas de fila (texto rotado 90°) ─
        if roi_h > roi_w * 1.5:
            gray_roi = cv2.rotate(gray_roi, cv2.ROTATE_90_CLOCKWISE)
            cell_h   = gray_roi.shape[0]
        else:
            cell_h   = roi_h

        # ── Mejora 2: Zoom adaptativo ────────────────────────────────────────
        if cell_h < 20:
            zoom = 4
        elif cell_h < 30:
            zoom = 3
        elif cell_h < 45:
            zoom = 3
        else:
            zoom = 2

        proc = clahe.apply(
            cv2.resize(gray_roi, None, fx=zoom, fy=zoom,
                       interpolation=cv2.INTER_CUBIC)
        )

        try:
            res = paddle_engine.ocr(proc, cls=True)
            text = ""
            if res and res[0]:
                # ── Mejora 1: Filtro por confianza mínima 0.85 ──────────────
                textos = [
                    line[1][0] for line in res[0]
                    if float(line[1][1]) >= CONF_THRESHOLD_V2
                ]
                if textos:
                    raw  = " ".join(textos)
                    # ── Mejora 4: Vocabulario arqueológico ──────────────────
                    text = _post_process_v2(clean_ocr_text(raw))

            final_cells.append(text)

        except Exception as e:
            print(f"  [OCR-batch error] celda {i}: {e}")
            final_cells.append("")

    print(f"  [OCR-Batch v2] {len(boxes)} celdas  conf>={CONF_THRESHOLD_V2}  zoom_adaptativo=ON  rotacion=ON")
    return final_cells




def recognize_cell_text(cell_img, cell_id: str) -> str:
    """
    Pipeline de OCR de alta precision para una celda:
    - Detecta automaticamente si la celda tiene texto vertical (h > 1.5*w)
    - Para texto vertical: prioriza variantes rotadas 90deg en el ensemble
    - Para texto horizontal: prioriza la version normal
    - Padding 10px + upscale adaptativo + ensemble 5 variantes
    """
    if get_paddle_ocr() is None or cell_img is None or cell_img.size == 0:
        return ""

    # ---- Determinar orientacion de la celda ----
    raw_h, raw_w = cell_img.shape[:2]
    is_portrait_cell = raw_h > raw_w * 1.5   # celda mas alta que ancha = texto vertical

    # ---- Padding forzado de 10px en todos los bordes ----
    PAD = 10
    if len(cell_img.shape) == 3:
        cell_padded = cv2.copyMakeBorder(cell_img, PAD, PAD, PAD, PAD,
                                          cv2.BORDER_CONSTANT, value=(255, 255, 255))
    else:
        cell_padded = cv2.copyMakeBorder(cell_img, PAD, PAD, PAD, PAD,
                                          cv2.BORDER_CONSTANT, value=255)

    # ---- Upscale adaptativo ----
    ch, cw = cell_padded.shape[:2]
    if min(ch, cw) < 64:
        cell_padded = cv2.resize(cell_padded, (cw * 4, ch * 4),
                                  interpolation=cv2.INTER_LANCZOS4)
    elif min(ch, cw) < 128:
        cell_padded = cv2.resize(cell_padded, (cw * 2, ch * 2),
                                  interpolation=cv2.INTER_LANCZOS4)

    # ---- Preprocesamiento: seleccion por varianza ----
    processed = preprocess_cell(cell_padded)

    # ---- Ensemble de 5 variantes ----
    rot180   = cv2.rotate(processed, cv2.ROTATE_180)
    rot90cw  = cv2.rotate(processed, cv2.ROTATE_90_CLOCKWISE)
    rot90ccw = cv2.rotate(processed, cv2.ROTATE_90_COUNTERCLOCKWISE)
    try:
        bilateral = cv2.bilateralFilter(
            processed if len(processed.shape) == 3 else cv2.cvtColor(processed, cv2.COLOR_GRAY2BGR),
            9, 75, 75
        )
        if len(bilateral.shape) == 3:
            bilateral = cv2.cvtColor(bilateral, cv2.COLOR_BGR2GRAY)
    except Exception:
        bilateral = processed

    variants = [
        (processed,  "pro"),
        (rot180,     "r180"),
        (rot90cw,    "r90cw"),
        (rot90ccw,   "r90ccw"),
        (bilateral,  "bilat"),
        (cell_padded, "raw"),  # Variante SIN preprocesamiento agresivo
    ]
    candidates = []

    for v_img, suffix in variants:
        # Forzamos det=False para las celdas individuales
        txt, conf = _run_ocr_on_image_with_conf(v_img, f"{cell_id}_{suffix}", det=False)
        if txt:
            candidates.append((txt, conf, suffix))

    if not candidates:
        return ""

    # ---- Seleccion inteligente segun orientacion ----
    if is_portrait_cell:
        # La celda es mas alta que ancha: el texto esta rotado 90deg
        # Preferir variantes rotadas (v2=90CW, v3=90CCW) multiplicando su score
        ROTATION_BOOST = 1.4
        scored = []
        for txt, conf, suffix in candidates:
            boost = ROTATION_BOOST if suffix in ('v2', 'v3') else 1.0
            scored.append((txt, conf * boost, suffix))
        best_candidate = max(scored, key=lambda x: x[1])
    else:
        # Celda normal horizontal: confiar en confianza pura
        scored = [(txt, conf, s) for txt, conf, s in candidates]
        best_candidate = max(scored, key=lambda x: x[1])

    best_text = best_candidate[0]

    # ---- Limpieza ----
    cleaned = clean_ocr_text(best_text)
    return postprocess_cell_text(cleaned)


# --------------------------------------------------------------------------
# ORDENAMIENTO DE CONTORNOS EN FILAS  (de ocr.py, mejorado)
# --------------------------------------------------------------------------

def sort_contours_to_rows(boxes):
    """
    Sincronizado con ocr2.py: Threshold dinámico basado en la altura promedio.
    """
    if not boxes: return []
    
    # Normalizar boxes si vienen con basura
    clean_boxes = []
    for b in boxes:
        if isinstance(b, tuple) and len(b) == 4: clean_boxes.append(b)
        elif isinstance(b, list) and len(b) > 0 and isinstance(b[0], tuple): clean_boxes.append(b[0])
        elif not isinstance(b, tuple): # Probablemente un contorno de cv2
            clean_boxes.append(cv2.boundingRect(b))
            
    # Ordenar por Y
    clean_boxes.sort(key=lambda x: x[1])
    
    rows = []
    curr = []
    avg_h = np.mean([b[3] for b in clean_boxes]) if clean_boxes else 20
    prev_y = None
    
    for b in clean_boxes:
        if prev_y is None or abs(b[1] - prev_y) <= avg_h * 0.5:
            curr.append(b)
        else:
            curr.sort(key=lambda x: x[0])
            # En el formato de matrix del backend, cada item es ((x,y,w,h), metadata)
            rows.append([(box, None) for box in curr])
            curr = [b]
        prev_y = b[1]
        
    if curr:
        curr.sort(key=lambda x: x[0])
        rows.append([(box, None) for box in curr])
        
    print(f"  [Sort] Reconstruidas {len(rows)} filas con threshold dinámico de {avg_h:.1f}px.")
    return rows


def reconstruct_table(roi, page_id: str, roi_original=None, precomputed_boxes=None, full_width=None):
    """
    Pipeline de reconstruccion de tabla de alta fidelidad.
    roi         = imagen PROCESADA (Limpia) - usada SOLO para detectar la estructura
    roi_original = imagen ORIGINAL sin procesar     - usada para OCR de texto de cada celda
    precomputed_boxes = (opcional) cajas detectadas externamente (vía ocr.py morph strategy)
    full_width  = ancho de la pagina completa para escalar kernels correctamente
    """
    steps = []
    H, W  = roi.shape[:2]
    # Si no se pasa full_width, usamos el de la ROI (comportamiento fallback)
    W_page = full_width if full_width else W
    global_page_keywords = []

    if precomputed_boxes is not None:
        # Si venimos del fallback morfologico, ya tenemos las cajas
        # Solo necesitamos agruparlas en filas
        print(f"  [Reconstruct] Usando {len(precomputed_boxes)} celdas precomputadas (Morfologia).")
        sorted_rows = sort_contours_to_rows(precomputed_boxes)
    else:
        # --- PASO 1: Escala de grises ---
        gray = cv2.cvtColor(roi, cv2.COLOR_BGR2GRAY)
        
        # --- PASO 2: Binarización (ocr2.py style) ---
        bin_img = cv2.threshold(gray, 128, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
        bin_img = cv2.bitwise_not(bin_img)
        
        # --- PASO 3: Estructura (REGLA ocr2.py: Escala con Ancho de ROI para celdas) ---
        # Si W_page es muy diferente a W, escalamos el kernel
        kl_v = max(W // 120, 3)
        kl_h = max(W // 40, 3)
        kern_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, kl_v))
        kern_h = cv2.getStructuringElement(cv2.MORPH_RECT, (kl_h, 1))
        
        v_lines = cv2.dilate(cv2.erode(bin_img, kern_v, iterations=3), kern_v, iterations=3)
        h_lines = cv2.dilate(cv2.erode(bin_img, kern_h, iterations=3), kern_h, iterations=3)
        
        # Combinación y limpieza (ocr2.py)
        grid = cv2.addWeighted(v_lines, 0.5, h_lines, 0.5, 0.0)
        grid = cv2.bitwise_not(grid)
        grid = cv2.erode(grid, cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3)), iterations=2)
        _, grid = cv2.threshold(grid, 0, 255, cv2.THRESH_OTSU)
        
        # --- PASO 4: Contornos con Filtros ocr2.py ---
        contours, _ = cv2.findContours(grid, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
        # Filtro area > 100 (ocr2.py) para no perder celdas pequeñas
        valid_boxes = []
        for cnt in contours:
            x, y, w, h = cv2.boundingRect(cnt)
            if cv2.contourArea(cnt) > 100 and 5 < h < 200:
                valid_boxes.append((x, y, w, h))
        
        sorted_rows = sort_contours_to_rows(valid_boxes)

    if not sorted_rows: return [], steps, ""
    
    # Visual step 6: Matrix indices
    contour_vis = roi.copy()
    for r_idx, row in enumerate(sorted_rows):
        for c_idx, ((cx, cy, cw, ch), _) in enumerate(row):
            cv2.rectangle(contour_vis, (cx, cy), (cx+cw, cy+ch), (0, 255, 0), 1)
            cv2.putText(contour_vis, f"{r_idx},{c_idx}", (cx+1, cy+10), cv2.FONT_HERSHEY_SIMPLEX, 0.3, (255,0,0), 1)
    steps.append(save_step(contour_vis, page_id, 6))

    # 3. Procesar OCR
    num_rows = len(sorted_rows)
    # Encontrar la fila con más columnas (usualmente los datos) para usarla como guía
    max_row_idx = 0
    # --- REGLA 4: RECONSTRUCION RULE-BASED ---
    num_rows = len(sorted_rows)
    if num_rows == 0: return [], steps, steps[-1] if steps else None
    
    # 1. Determinar ancho maximo (num_cols)
    num_cols = 0
    max_row_idx = 0
    for r_idx, row in enumerate(sorted_rows):
        if len(row) > num_cols:
            num_cols = len(row)
            max_row_idx = r_idx
    
    if num_cols == 0: return [], steps, steps[-1] if steps else None
    
    # 2. Dividir Cabeceras Fusionadas (Regla 4.2)
    # Si la fila 0 tiene Pocas celdas comparada con el maximo, dividimos usando la guia
    if num_rows > 0 and len(sorted_rows[0]) < num_cols * 0.7:
        print(f"  [Reconstruct] Dividiendo Cabecera Merged (Regla 4.2) usando Fila {max_row_idx}")
        guia = sorted_rows[max_row_idx]
        hy_min = sorted_rows[0][0][0][1]
        hy_max = hy_min + sorted_rows[0][0][0][3]
        new_header = []
        for g_box, _ in guia:
            gx, _, gw, _ = g_box
            new_header.append(((gx, hy_min, gw, hy_max - hy_min), None))
        sorted_rows[0] = new_header

    # 3. Búsqueda de Cabeceras (Regla 4.1) - STRIP OCR (Mejora Crítica)
    target_img = roi_original if roi_original is not None else roi
    header_found = False
    header_idx = 0
    final_matrix = []
    
    # Escaneamos primeras 15 filas (Page 3 tiene cabeceras complejas)
    for r in range(min(15, num_rows)):
        row_boxes = [c[0] for c in sorted_rows[r]]
        if not row_boxes: continue
        
        # Calculamos el BBox de la fila completa para Strip OCR
        rx1 = min(b[0] for b in row_boxes); ry1 = min(b[1] for b in row_boxes)
        rx2 = max(b[0]+b[2] for b in row_boxes); ry2 = max(b[1]+b[3] for b in row_boxes)
        # Seccion de la imagen original
        row_strip = target_img[max(0, ry1-2):min(H, ry2+2), max(0, rx1-2):min(W, rx2+2)]
        
        if row_strip is None or row_strip.size == 0: continue
        
        for deg in [-90, 90, 0]:
            # Rotamos la tira completa segun el angulo
            if deg == -90:
                rot_strip = cv2.rotate(row_strip, cv2.ROTATE_90_CLOCKWISE)
            elif deg == 90:
                rot_strip = cv2.rotate(row_strip, cv2.ROTATE_90_COUNTERCLOCKWISE)
            else:
                rot_strip = row_strip

            # OCR en la tira completa (Paddle es mejor para strips)
            res = get_paddle_ocr().ocr(rot_strip, cls=False)
            if not res or not res[0]: continue
            
            strip_text = " ".join([line[1][0].upper() for line in res[0]])
            # Busqueda de al menos 2 palabras clave Arqueologicas
            matches = [kw for kw in ARCHAEO_VOCAB if len(kw) > 3 and kw in strip_text]
            
            # Fallback a difflib por palabra en el strip detectado
            if len(matches) < 2:
                words = strip_text.split()
                for w in words:
                    if len(w) > 3:
                        m = difflib.get_close_matches(w, ARCHAEO_VOCAB, n=1, cutoff=0.6)
                        if m: matches.append(m[0])
            
            if len(set(matches)) >= 2:
                header_idx = r
                header_found = True
                print(f"  [Reconstruct] Header CONFIRMADA en fila {r} (Orientacion {deg}). KWs: {list(set(matches))[:4]}")
                
                # Para la matriz final, usamos el reconocimiento por celda con la orientacion ganadora
                res_texts = recognize_cells_batch(target_img, row_boxes, rotate_deg=deg, margin=5)
                row_clean = [postprocess_cell_text(t) for t in res_texts]
                header_row = (row_clean + [""] * num_cols)[:num_cols]
                final_matrix = [header_row]
                # Guardamos keywords para el reporte
                global_page_keywords = list(set(matches))
                break
        if header_found: break

    # Fallback si no hay cabecera (Regla 4.3)
    if not header_found:
        print("  [WARNING] No se detecto cabecera formal. Usando Fila 0 como cabecera por defecto.")
        row_boxes = [c[0] for c in sorted_rows[0]]
        res_texts = recognize_cells_batch(target_img, row_boxes, rotate_deg=0, margin=5)
        header_row = ( [postprocess_cell_text(t) for t in res_texts] + [""] * num_cols )[:num_cols]
        final_matrix = [header_row]
        global_page_keywords = []

    # 4. Procesar Datos (Batch Optimizado)
    data_start = header_idx + 1
    all_data_boxes = []
    box_to_pos = [] 
    
    for r in range(data_start, num_rows):
        for c_idx, (box, _) in enumerate(sorted_rows[r]):
            if c_idx < num_cols:
                all_data_boxes.append(box)
                box_to_pos.append((r, c_idx))
                
    if all_data_boxes:
        print(f"  [Reconstruct] Procesando {len(all_data_boxes)} celdas de datos en batch (shave=2)...")
        all_res_texts = recognize_cells_batch(target_img, all_data_boxes, rotate_deg=0, allowlist=DATA_ALLOWLIST, margin=5)
        
        # Mapear de vuelta a la matriz
        # Inicializamos matriz de datos (vacia)
        data_matrix = [["" for _ in range(num_cols)] for _ in range(num_rows - data_start)]
        
        for i, raw_t in enumerate(all_res_texts):
            r_idx_orig, c_idx = box_to_pos[i]
            r_idx = r_idx_orig - data_start
            
            clean_t = postprocess_cell_text(raw_t)
            
            # Refuerzo para celdas numéricas (Regla 3.1)
            # Solo si no es un texto largo arqueológico (vocabulario)
            is_keyword = any(kw in clean_t.upper() for kw in ARCHAEO_VOCAB)
            if not is_keyword and len(clean_t) < 8 and any(ch.isdigit() for ch in clean_t):
                digits = "".join([ch for ch in clean_t if ch.isdigit()])
                if digits: clean_t = digits
            
            data_matrix[r_idx][c_idx] = clean_t
            
        final_matrix.extend(data_matrix)
    
    # Reportar keywords detectados
    res_info = f"Keywords: {', '.join(global_page_keywords)}" if global_page_keywords else "No keywords"

    # 5. Visualizacion Final
    final_vis = roi.copy()
    for r_idx, row_boxes in enumerate(sorted_rows):
        for box_tuple in row_boxes:
            b = box_tuple[0]
            cv2.rectangle(final_vis, (b[0], b[1]), (b[0]+b[2], b[1]+b[3]), (0, 255, 0), 1)
    
    steps.append(save_step(final_vis, page_id, 7))
    print(f"  [Reconstruct] Tabla terminada: {len(final_matrix)} filas, {num_cols} columnas.")
    
    return compact_matrix(final_matrix), steps, steps[-1], global_page_keywords


def compact_matrix(matrix):
    """Elimina filas y columnas donde TODOS los valores están vacíos."""
    if not matrix:
        return matrix

    # Normalizar: todas las filas al mismo ancho
    max_cols = max(len(r) for r in matrix)
    padded   = [r + [""] * (max_cols - len(r)) for r in matrix]

    # Columnas con al menos un valor no vacío
    kept_cols = [c for c in range(max_cols)
                 if any(padded[r][c].strip() for r in range(len(padded)))]

    # Filtrar columnas
    filtered = [[padded[r][c] for c in kept_cols] for r in range(len(padded))]

    # Filas con al menos un valor no vacío
    result = [row for row in filtered if any(v.strip() for v in row)]

    return result if result else matrix  # Devolver original si nada queda


def process_image(img_path: str, page_num: int) -> dict:
    """
    Pipeline OCR fiel a ocr2.py: imagen completa, kernels basados en W real,
    paso directo de numpy a PaddleOCR. Maximo fidelidad con el script de referencia.
    """
    print(f"\n[PAGE {page_num}] {os.path.basename(img_path)}")
    if not os.path.exists(img_path):
        return {"error": "Archivo no encontrado"}

    # Leer imagen robusta (imdecode con fromfile -> funciona con rutas Unicode Windows)
    img = imread_unicode(img_path)
    if img is None:
        return {"error": "No se pudo decodificar"}

    orig_h, orig_w = img.shape[:2]
    page_id = str(page_num)
    print(f"  [IMG] {orig_w}x{orig_h} px")

    steps = []
    steps.append(save_step(img, page_id, 0))

    # === PASO 1: DETECCION DE GRILLA (replica exacta de ocr2.py) ===
    print("  [Grid] Detectando celdas sobre imagen completa (logica ocr2.py)...")
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    bin_img = cv2.bitwise_not(
        cv2.threshold(gray, 128, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
    )

    # Kernels basados en W de la imagen COMPLETA (clave de ocr2.py)
    kl_v = max(orig_w // 120, 3)
    kl_h = max(orig_w // 40, 3)
    v_k = cv2.getStructuringElement(cv2.MORPH_RECT, (1, kl_v))
    h_k = cv2.getStructuringElement(cv2.MORPH_RECT, (kl_h, 1))

    v_l = cv2.dilate(cv2.erode(bin_img, v_k, iterations=3), v_k, iterations=3)
    h_l = cv2.dilate(cv2.erode(bin_img, h_k, iterations=3), h_k, iterations=3)

    grid = cv2.addWeighted(v_l, 0.5, h_l, 0.5, 0.0)
    grid = cv2.bitwise_not(grid)
    grid = cv2.erode(grid, cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3)), iterations=2)
    _, grid = cv2.threshold(grid, 0, 255, cv2.THRESH_OTSU)

    cnts, _ = cv2.findContours(grid, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
    boxes = []
    for c in cnts:
        x, y, w, h = cv2.boundingRect(c)
        if cv2.contourArea(c) > 100 and 5 < h < 200:
            boxes.append((x, y, w, h))

    print(f"  [Grid] Kernels V=1x{kl_v}, H={kl_h}x1 -> {len(boxes)} celdas")

    if len(boxes) < 5:
        print("  [Grid] Pocas celdas, intentando con AdaptiveThreshold...")
        img_gray_red = img[:, :, 2]
        img_bin_alt = cv2.adaptiveThreshold(img_gray_red, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 15, 10)
        v_l2 = cv2.dilate(cv2.erode(img_bin_alt, v_k, iterations=3), v_k, iterations=3)
        h_l2 = cv2.dilate(cv2.erode(img_bin_alt, h_k, iterations=3), h_k, iterations=3)
        grid2 = cv2.addWeighted(v_l2, 0.5, h_l2, 0.5, 0.0)
        grid2 = cv2.erode(cv2.bitwise_not(grid2), cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3)), iterations=2)
        _, grid2 = cv2.threshold(grid2, 0, 255, cv2.THRESH_OTSU)
        cnts2, _ = cv2.findContours(grid2, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
        boxes = []
        for c in cnts2:
            x, y, w, h = cv2.boundingRect(c)
            if cv2.contourArea(c) > 400 and 10 < h < 250:
                boxes.append((x, y, w, h))
        print(f"  [Grid-Alt] {len(boxes)} celdas con AdaptiveThreshold")

    # === PASO 2: ORDENAR CELDAS EN FILAS (replica ocr2.py) ===
    if not boxes:
        print("  [WARN] No se detectaron celdas.")
        return {"page": page_num, "regions": [], "visual_steps": steps}

    boxes.sort(key=lambda b: b[1])
    rows_ocr2 = []
    curr_row = []
    avg_h = np.mean([b[3] for b in boxes]) if boxes else 20
    prev_y = None

    for b in boxes:
        if prev_y is None or abs(b[1] - prev_y) <= avg_h * 0.5:
            curr_row.append(b)
        else:
            curr_row.sort(key=lambda b: b[0])
            rows_ocr2.append(curr_row)
            curr_row = [b]
        prev_y = b[1]
    if curr_row:
        curr_row.sort(key=lambda b: b[0])
        rows_ocr2.append(curr_row)

    num_cols = max(len(r) for r in rows_ocr2) if rows_ocr2 else 0
    print(f"  [Sort] {len(rows_ocr2)} filas x {num_cols} columnas (avg_h={avg_h:.1f})")

    # Visualizacion de la grilla detectada
    vis = img.copy()
    for row in rows_ocr2:
        for (x, y, w, h) in row:
            cv2.rectangle(vis, (x, y), (x+w, y+h), (0, 200, 0), 1)
    steps.append(save_step(vis, page_id, 6))

    # === PASO 3: OCR POR CELDA (replica ocr2.py) ===
    print(f"  [OCR] Procesando {sum(len(r) for r in rows_ocr2)} celdas con GPU...")
    paddle_engine = get_paddle_ocr()
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    MARGIN = 5

    matrix = [["" for _ in range(num_cols)] for _ in range(len(rows_ocr2))]

    for r_i, row in enumerate(rows_ocr2):
        for c_i, (x, y, w, h) in enumerate(row):
            if c_i >= num_cols:
                continue
            # Recorte con margen externo (ocr2.py)
            iy = max(y - MARGIN, 0)
            ix = max(x - MARGIN, 0)
            roi = img[iy:min(iy+h+2*MARGIN, orig_h), ix:min(ix+w+2*MARGIN, orig_w)]
            if roi.size == 0:
                continue
            # Preprocesamiento ocr2.py: Gray -> Resize x2 -> CLAHE
            roi_gray = cv2.cvtColor(roi, cv2.COLOR_BGR2GRAY)
            roi_proc = cv2.resize(roi_gray, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
            roi_proc = clahe.apply(roi_proc)
            try:
                res = paddle_engine.ocr(roi_proc, cls=True)
                if res and res[0]:
                    text = " ".join([line[1][0] for line in res[0]])
                    matrix[r_i][c_i] = clean_ocr_text(text)
            except Exception as e:
                pass  # celda vacia

    print(f"  [OCR] Completado. Matriz: {len(matrix)} filas x {num_cols} columnas")

    # === PASO 4: COMPACTAR MATRIZ (igual que ocr2.py) ===
    # NO usar compact_matrix aqui - conservar estructura original
    header = matrix[0] if matrix else []
    data_rows = matrix[1:] if len(matrix) > 1 else []

    # Preparar resultado como region de tabla
    result_region = {
        "id": f"region_{page_num}_0",
        "type": "table",
        "bbox": [0, 0, orig_w, orig_h],
        "content": {
            "header": header,
            "rows": data_rows,
            "keywords": []
        },
        "result_image_url": steps[-1] if steps else ""
    }

    return {
        "page": page_num,
        "regions": [result_region],
        "visual_steps": steps
    }


# --------------------------------------------------------------------------
# ENDPOINTS
# --------------------------------------------------------------------------

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(('.pdf', '.png', '.jpg', '.jpeg')):
        raise HTTPException(status_code=400, detail="Use PDF, PNG o JPG.")

    req_id = str(uuid.uuid4())
    
    # 1. Guardar archivo temporalmente
    temp_path = os.path.join(UPLOADS_DIR, f"temp_{req_id}_{file.filename}")
    with open(temp_path, "wb") as buf:
        shutil.copyfileobj(file.file, buf)

    # 2. Calcular Hash y revisar Caché (Opcional - desactivado para ver cambios en vivo)
    md5_hash = file_hash(temp_path)
    # cached_data = get_cached(md5_hash)
    # if cached_data:
    #     os.remove(temp_path)
    #     update_progress(req_id, 100, "Cargado de caché", done=True)
    #     return JSONResponse(content={"request_id": req_id, "pages": cached_data})

    # Renombrar archivo a su hash final
    file_path = os.path.join(UPLOADS_DIR, f"{md5_hash}_{file.filename}")
    if os.path.exists(file_path):
        os.remove(temp_path) # Ya existía
    else:
        os.rename(temp_path, file_path)

    print(f"\n[UPLOAD] {file.filename} -> {file_path}")
    update_progress(req_id, 10, "Convirtiendo...", done=False)

    try:
        results = []
        if file.filename.lower().endswith(".pdf"):
            images = convert_from_path(file_path, dpi=200, poppler_path=POPPLER_DIR, last_page=10)
            
            # Guardar todas las imgs primero
            page_tasks = []
            for i, pil_img in enumerate(images):
                page_path = os.path.join(UPLOADS_DIR, f"{md5_hash}_p{i}.jpg")
                pil_img.save(page_path, "JPEG")
                
                # Función asíncrona wrapper para correr process_image en thread
                async def run_page(p_path, p_num):
                    update_progress(req_id, 20 + int(70*(p_num/len(images))), f"Procesando pág. {p_num}/{len(images)}")
                    loop = asyncio.get_event_loop()
                    # process_image es sincrona, enviar a ThreadPoolExecutor
                    data = await loop.run_in_executor(None, process_image, p_path, p_num)
                    data["image_url"] = f"/uploads/{os.path.basename(p_path)}"
                    return data
                
                page_tasks.append(run_page(page_path, i + 1))
            
            # Ejecutar todas las páginas en paralelo
            results = await asyncio.gather(*page_tasks)
            # Ordenar por número de página
            results = sorted(results, key=lambda x: x["page"])

        else:
            # Procesar imagen única
            update_progress(req_id, 50, "Procesando imagen (GPU)...")
            loop = asyncio.get_event_loop()
            page_data = await loop.run_in_executor(None, process_image, file_path, 1)
            page_data["image_url"] = f"/uploads/{os.path.basename(file_path)}"
            results.append(page_data)

        # Guardar en Caché
        set_cached(md5_hash, results)
        
        update_progress(req_id, 100, "¡Completado!", done=True)
        return JSONResponse(content={"request_id": req_id, "pages": results})

    except Exception as e:
        import traceback
        traceback.print_exc()
        update_progress(req_id, 0, f"Error: {e}", done=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/export")
async def export_excel(data: dict):
    req_id   = data.get("request_id", str(uuid.uuid4()))
    pages    = data.get("pages", [])
    out_path = os.path.join(OUTPUT_DIR, f"resultado_{req_id}.xlsx")

    table_count = 0
    try:
        with pd.ExcelWriter(out_path, engine='openpyxl') as writer:
            for page in pages:
                for i, region in enumerate(page.get("regions", [])):
                    if region.get("type") != "table" or not region.get("content"):
                        continue
                    header = region["content"].get("header", [])
                    rows   = region["content"].get("rows", [])
                    df     = pd.DataFrame(rows, columns=header if header else None)
                    sheet  = f"P{page['page']}_T{i+1}"[:31]
                    df.to_excel(writer, sheet_name=sheet, index=False)
                    table_count += 1

            if table_count == 0:
                pd.DataFrame({"Info": ["No se detectaron tablas."]}).to_excel(
                    writer, sheet_name="Sin_Tablas", index=False)

        return FileResponse(out_path,
                            filename=f"arcaica_ocr_{req_id}.xlsx",
                            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error Excel: {e}")


# --------------------------------------------------------------------------
# SSE PROGRESS ENDPOINT
# --------------------------------------------------------------------------

@app.get("/progress/{job_id}")
async def sse_progress(job_id: str):
    """Server-Sent Events: transmite progreso en tiempo real al frontend."""
    async def event_stream():
        timeout = 300   # máximo 5 min esperando
        start   = time.time()
        last_pct = -1

        while time.time() - start < timeout:
            state = JOB_PROGRESS.get(job_id)
            if state is None:
                yield f"data: {json.dumps({'pct': 0, 'msg': 'Esperando...', 'done': False})}\n\n"
                await asyncio.sleep(0.5)
                continue

            pct  = state.get("pct", 0)
            msg  = state.get("msg", "")
            done = state.get("done", False)

            if pct != last_pct or done:
                last_pct = pct
                yield f"data: {json.dumps({'pct': pct, 'msg': msg, 'done': done})}\n\n"

            if done:
                break
            await asyncio.sleep(0.3)

    return StreamingResponse(event_stream(),
                             media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache",
                                      "X-Accel-Buffering": "no"})


# --------------------------------------------------------------------------
# EXPORT CSV
# --------------------------------------------------------------------------

@app.post("/export/csv")
async def export_csv(data: dict):
    """Exporta todas las tablas a un único CSV concatenado."""
    pages    = data.get("pages", [])
    req_id   = data.get("request_id", str(uuid.uuid4()))
    out_path = os.path.join(OUTPUT_DIR, f"resultado_{req_id}.csv")

    rows_all = []
    for page in pages:
        for i, region in enumerate(page.get("regions", [])):
            if region.get("type") != "table" or not region.get("content"):
                continue
            header = region["content"].get("header", [])
            rows   = region["content"].get("rows", [])
            # Cabecera de sección
            rows_all.append([f"=== Página {page['page']} - Tabla {i+1} ==="] + [""] * max(0, len(header)-1))
            if header:
                rows_all.append(header)
            rows_all.extend(rows)
            rows_all.append([])  # línea vacía separadora

    if not rows_all:
        rows_all = [["No se detectaron tablas."]]

    df = pd.DataFrame(rows_all)
    df.to_csv(out_path, index=False, header=False, encoding="utf-8-sig")

    return FileResponse(out_path,
                        filename=f"arcaica_ocr_{req_id}.csv",
                        media_type="text/csv")


# --------------------------------------------------------------------------
# EXPORT WORD (.docx)
# --------------------------------------------------------------------------

@app.post("/export/word")
async def export_word(data: dict):
    """Exporta tablas a documento Word (.docx) con formato."""
    if not HAS_DOCX:
        raise HTTPException(status_code=501, detail="python-docx no instalado")

    pages    = data.get("pages", [])
    req_id   = data.get("request_id", str(uuid.uuid4()))
    out_path = os.path.join(OUTPUT_DIR, f"resultado_{req_id}.docx")

    doc = DocxDocument()
    doc.add_heading("Resultados OCR - ArcheoOCR", 0)

    table_count = 0
    for page in pages:
        doc.add_heading(f"Página {page.get('page', '?')}", level=1)

        for i, region in enumerate(page.get("regions", [])):
            rtype = region.get("type", "")

            if rtype == "text":
                content = region.get("content", {})
                text    = content.get("text", "") if isinstance(content, dict) else str(content)
                if text.strip():
                    doc.add_paragraph(text)

            elif rtype == "table":
                content = region.get("content")
                if not content:
                    continue
                header = content.get("header", [])
                rows   = content.get("rows", [])
                all_rows = ([header] if header else []) + rows
                if not all_rows:
                    continue

                doc.add_heading(f"Tabla {i+1}", level=2)
                n_cols = max(len(r) for r in all_rows)
                tbl    = doc.add_table(rows=len(all_rows), cols=n_cols)
                tbl.style = "Table Grid"

                for r_idx, row_data in enumerate(all_rows):
                    for c_idx, cell_val in enumerate(row_data):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = str(cell_val) if cell_val else ""
                        if r_idx == 0 and header:
                            # Negrita para encabezado
                            for run in cell.paragraphs[0].runs:
                                run.bold = True

                table_count += 1
                doc.add_paragraph()  # espacio

    if table_count == 0:
        doc.add_paragraph("No se detectaron tablas en el documento.")

    doc.save(out_path)
    return FileResponse(
        out_path,
        filename=f"arcaica_ocr_{req_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# --------------------------------------------------------------------------
# EXPORT PDF (texto plano de tablas)
# --------------------------------------------------------------------------

@app.post("/export/pdf")
async def export_pdf_report(data: dict):
    """Exporta tablas a PDF con fpdf2."""
    if not HAS_FPDF:
        raise HTTPException(status_code=501, detail="fpdf2 no instalado")

    pages    = data.get("pages", [])
    req_id   = data.get("request_id", str(uuid.uuid4()))
    out_path = os.path.join(OUTPUT_DIR, f"resultado_{req_id}.pdf")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Resultados OCR - ArcheoOCR", ln=True, align="C")
    pdf.ln(5)

    for page in pages:
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, f"Pagina {page.get('page', '?')}", ln=True)
        pdf.ln(2)

        for i, region in enumerate(page.get("regions", [])):
            rtype   = region.get("type", "")
            content = region.get("content", {})

            if rtype == "text":
                text = content.get("text", "") if isinstance(content, dict) else str(content)
                pdf.set_font("Helvetica", "", 10)
                pdf.multi_cell(0, 6, text[:2000])
                pdf.ln(3)

            elif rtype == "table":
                if not content:
                    continue
                header = content.get("header", [])
                rows   = content.get("rows", [])
                all_r  = ([header] if header else []) + rows
                if not all_r:
                    continue

                pdf.set_font("Helvetica", "B", 11)
                pdf.cell(0, 7, f"Tabla {i+1}", ln=True)
                
                # Render simple table
                n_cols  = max(len(r) for r in all_r)
                if n_cols == 0: continue
                
                # FPDF2 table feature
                pdf.set_font("Helvetica", "", 8)
                with pdf.table() as table:
                    for r_idx, row_data in enumerate(all_r[:40]): # Limitar a 40 para que no desborde excesivamente
                        row = table.row()
                        for c_idx in range(n_cols):
                            val = str(row_data[c_idx]) if c_idx < len(row_data) else ""
                            # Truncate text loosely to fit in PDF
                            row.cell(val[:35])
                pdf.ln(4)

    pdf.output(out_path)
    return FileResponse(out_path,
                        filename=f"arcaica_ocr_{req_id}.pdf",
                        media_type="application/pdf")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
